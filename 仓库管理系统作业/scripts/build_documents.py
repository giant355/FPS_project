from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = r"C:\Users\Lenovo\Desktop\新建文件夹\My project (11)\仓库管理系统作业"


def set_font(run, name="宋体", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_font(r, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn("w:" + name))
        if border is None:
            border = OxmlElement("w:" + name)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "D9E2EC")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
        if widths:
            table.rows[0].cells[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
            if widths:
                cells[i].width = widths[i]
    set_table_borders(table)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 14, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
    return doc


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size=12)
    p.paragraph_format.first_line_indent = Pt(24) if style is None else None
    return p


def add_cover(doc, title):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("数据库原理与应用实验")
    set_font(r, size=22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("综合作品")
    set_font(r, size=22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("2025-2026学年第2学期")
    set_font(r, size=14)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, size=20, bold=True)
    for _ in range(4):
        doc.add_paragraph()
    info = [
        "题目：仓库零件管理系统",
        "组长：魏皓荣  2024113060524",
        "组员：张珈宸  2023116000227",
        "数据库：MySQL 5.5",
    ]
    for item in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        set_font(r, size=13)
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_paper():
    doc = setup_doc()
    add_cover(doc, "仓库零件管理系统的设计与实现")

    doc.add_heading("仓库零件管理系统的设计与实现", level=1)
    doc.add_heading("1 前言", level=1)
    doc.add_heading("1.1 研究背景", level=2)
    add_para(doc, "随着企业和学校实训室物资种类不断增加，传统手工登记方式在零件入库、修改、查找和删除等方面容易出现记录不统一、查询效率低、数据更新不及时等问题。仓库零件管理系统通过数据库保存零件号、零件名、零件颜色、零件数量和购买时间等信息，可以使仓库零件管理更加规范。")
    add_para(doc, "本作品以课程实验要求为基础，采用 MySQL 5.5 建立数据库，并使用简单的 PHP 页面实现对零件信息的增、删、改、查操作。系统规模适中，逻辑清晰，适合作为数据库原理与应用课程的综合实践。")
    doc.add_heading("1.2 研究意义", level=2)
    add_para(doc, "仓库零件信息具有数量变化频繁、查询需求明确、字段结构稳定等特点，适合使用关系型数据库进行管理。通过本系统可以减少纸质登记带来的遗漏和重复问题，提高零件查询速度，方便管理员掌握库存情况。")
    add_para(doc, "从学习角度看，本系统覆盖数据库创建、数据表设计、主码设置、数据插入以及通过程序语言实现 CRUD 功能等内容，能够帮助学生理解数据库设计与程序开发之间的联系。")
    doc.add_heading("1.3 研究内容", level=2)
    add_para(doc, "本文围绕仓库零件管理系统展开设计与实现，主要研究内容包括：分析系统需求；设计功能模块；完成数据库概念结构和逻辑结构设计；编写 MySQL 5.5 数据库脚本；使用 PHP 编写添加、修改、删除、查看和查询页面；最后对系统实现效果进行总结。")

    doc.add_heading("2 相关技术", level=1)
    doc.add_heading("2.1 MySQL 5.5 数据库", level=2)
    add_para(doc, "MySQL 是常用的关系型数据库管理系统，具有体积小、运行稳定、语法清晰等特点。MySQL 5.5 支持 InnoDB 存储引擎、主键约束、日期类型、字符集设置和常见 SQL 操作，能够满足本系统的数据存储需求。")
    doc.add_heading("2.2 PHP 与 MySQL 访问", level=2)
    add_para(doc, "PHP 是一种常见的服务器端脚本语言，适合开发简单的信息管理系统。本作品使用 mysqli 方式连接 MySQL 数据库，通过 SQL 语句完成零件信息的查询、添加、修改和删除。")
    doc.add_heading("2.3 HTML 与 CSS", level=2)
    add_para(doc, "HTML 用于构建系统页面结构，包括零件列表、录入表单和详情页面；CSS 用于设置页面布局、表格、按钮和提示信息，使系统界面更加清楚。")

    doc.add_heading("3 系统总体设计", level=1)
    doc.add_heading("3.1 需求分析", level=2)
    add_para(doc, "系统面向仓库管理员使用，管理对象为仓库中的零件信息。每条零件记录包括零件号、零件名、零件颜色、零件数量和购买时间。系统需要保证零件号唯一，并能够对库存数量进行非负校验。")
    add_para(doc, "系统基本功能包括：零件信息添加、零件信息修改、零件信息删除和查看功能。同时，为提高实用性，列表页面增加按零件号、零件名和颜色进行模糊查询的功能。")
    add_table(doc, ["需求编号", "功能名称", "功能说明"], [
        ["R1", "添加零件", "录入零件号、名称、颜色、数量、购买时间等信息"],
        ["R2", "修改零件", "根据零件号定位记录并修改名称、颜色、数量、购买时间等字段"],
        ["R3", "删除零件", "删除不再需要保存的零件记录"],
        ["R4", "查看零件", "展示零件列表和单条零件详情"],
        ["R5", "查询零件", "按零件号、零件名、颜色进行模糊查询"],
    ], [Cm(2.2), Cm(3.2), Cm(10.2)])

    doc.add_heading("3.2 功能模块设计", level=2)
    add_para(doc, "仓库零件管理系统采用简单的单用户管理结构，由数据库层、业务处理层和页面显示层组成。数据库层保存 parts 表；业务处理层负责执行 SQL 语句；页面显示层向用户提供表单、列表和详情信息。")
    add_table(doc, ["模块", "主要页面", "说明"], [
        ["零件列表与查询模块", "index.php", "显示全部零件记录，支持关键字查询"],
        ["添加模块", "part_add.php", "提交表单后向 parts 表插入新记录"],
        ["修改模块", "part_edit.php", "读取原记录并更新对应字段"],
        ["删除模块", "part_delete.php", "根据零件号删除指定记录"],
        ["查看模块", "part_view.php", "显示单条零件记录的完整信息"],
    ], [Cm(3.8), Cm(3.8), Cm(8.0)])

    doc.add_heading("3.3 数据库设计", level=2)
    doc.add_heading("3.3.1 概念结构设计", level=3)
    add_para(doc, "本系统的核心实体为“零件”。一个零件实体由零件号、零件名、零件颜色、零件数量和购买时间组成。零件号用于唯一标识零件记录，是该实体的主码。")
    add_table(doc, ["实体", "属性"], [
        ["零件", "零件号、零件名、零件颜色、零件数量、购买时间、备注、创建时间、修改时间"],
    ], [Cm(3.0), Cm(12.6)])
    add_para(doc, "概念结构可描述为：仓库管理员对零件实体进行维护，维护操作包括添加、修改、删除和查看；零件实体保存仓库管理所需的基础资料。")

    doc.add_heading("3.3.2 逻辑结构设计", level=3)
    add_para(doc, "数据库名称为 warehouse_db，核心数据表为 parts。该表以 part_no 作为主码，保证每个零件编号唯一；quantity 字段使用 INT 类型保存库存数量；purchase_time 字段使用 DATE 类型保存购买时间。")
    add_table(doc, ["字段名", "数据类型", "长度", "约束", "说明"], [
        ["part_no", "VARCHAR", "20", "PRIMARY KEY, NOT NULL", "零件号"],
        ["part_name", "VARCHAR", "50", "NOT NULL", "零件名"],
        ["part_color", "VARCHAR", "20", "NOT NULL", "零件颜色"],
        ["quantity", "INT", "-", "NOT NULL, DEFAULT 0", "零件数量"],
        ["purchase_time", "DATE", "-", "NOT NULL", "购买时间"],
        ["remark", "VARCHAR", "200", "NULL", "备注"],
        ["created_at", "DATETIME", "-", "NOT NULL", "创建时间"],
        ["updated_at", "DATETIME", "-", "NOT NULL", "修改时间"],
    ], [Cm(3.0), Cm(2.5), Cm(1.6), Cm(4.5), Cm(4.0)])

    doc.add_heading("4 系统的功能实现", level=1)
    doc.add_heading("4.1 创建数据库和数据表", level=2)
    add_para(doc, "在 database/warehouse.sql 中使用 CREATE DATABASE 创建 warehouse_db 数据库，并使用 CREATE TABLE 创建 parts 表。脚本设置 DEFAULT CHARSET=utf8，便于保存中文零件名称和颜色。")
    doc.add_heading("4.2 添加功能实现", level=2)
    add_para(doc, "添加页面 part_add.php 通过表单接收零件号、零件名、颜色、数量和购买时间等信息。程序先判断必填字段是否为空，并检查数量不能小于 0，再执行 INSERT INTO parts 语句写入数据库。")
    doc.add_heading("4.3 修改功能实现", level=2)
    add_para(doc, "修改页面 part_edit.php 根据零件号读取原有记录并填充到表单中。用户修改后提交，程序执行 UPDATE parts SET ... WHERE part_no=... 语句更新对应记录，同时修改 updated_at 字段。")
    doc.add_heading("4.4 删除功能实现", level=2)
    add_para(doc, "删除页面 part_delete.php 接收零件号参数，执行 DELETE FROM parts WHERE part_no=... 语句删除记录。列表页面在删除链接中加入确认提示，降低误删概率。")
    doc.add_heading("4.5 查看与查询功能实现", level=2)
    add_para(doc, "查看功能包括列表查看和详情查看。index.php 显示所有零件信息，并支持输入关键字对零件号、零件名和颜色进行模糊查询；part_view.php 用于显示单条零件记录的详细内容。")
    add_table(doc, ["操作", "主要 SQL 语句"], [
        ["添加", "INSERT INTO parts(...) VALUES(...)"],
        ["修改", "UPDATE parts SET ... WHERE part_no='编号'"],
        ["删除", "DELETE FROM parts WHERE part_no='编号'"],
        ["查看", "SELECT * FROM parts ORDER BY purchase_time DESC"],
        ["查询", "SELECT * FROM parts WHERE part_no LIKE ... OR part_name LIKE ..."],
    ], [Cm(3.0), Cm(12.6)])

    doc.add_heading("5 总结", level=1)
    add_para(doc, "本作品完成了仓库零件管理系统的基本设计与实现，能够对零件号、零件名、零件颜色、零件数量和购买时间等数据进行管理。系统数据库结构简单清晰，主码设置合理，功能覆盖课程要求中的添加、修改、删除和查看。")
    add_para(doc, "通过本次综合作品，进一步熟悉了 MySQL 数据库建库建表、数据类型选择、主键约束设置以及 SQL 增删改查语句的使用，也加深了对程序页面与数据库交互过程的理解。后续如果继续完善，可以增加用户登录、库存预警、导出报表和操作日志等功能。")
    doc.save(OUT_DIR + r"\论文-仓库零件管理系统的设计与实现.docx")


def build_assignment():
    doc = setup_doc()
    add_cover(doc, "仓库零件管理系统组员分工及完成情况")
    doc.add_heading("组员分工及完成情况说明书", level=1)
    add_para(doc, "本说明书用于说明仓库零件管理系统综合作品中两名组员的任务分配、完成内容和协作情况。")
    add_table(doc, ["姓名", "学号", "角色", "主要任务", "完成情况"], [
        ["魏皓荣", "2024113060524", "组长", "负责需求分析、数据库设计、SQL 脚本编写、论文整体整理与最终检查。", "已完成"],
        ["张珈宸", "2023116000227", "组员", "负责 PHP 功能页面编写、增删改查逻辑整理、界面样式和源文件说明。", "已完成"],
    ], [Cm(2.2), Cm(3.4), Cm(2.0), Cm(6.4), Cm(2.0)])
    doc.add_heading("一、任务分配说明", level=1)
    add_para(doc, "魏皓荣作为组长，主要负责确定系统题目和功能范围，完成仓库零件信息的数据项分析，设计数据库 warehouse_db 和 parts 数据表，并整理论文中前言、总体设计、数据库设计和总结等部分。")
    add_para(doc, "张珈宸主要负责根据数据库结构编写 PHP 页面，包括零件列表查询、添加、修改、删除和查看页面，同时整理 README 运行说明和页面样式文件。")
    doc.add_heading("二、完成情况", level=1)
    add_para(doc, "本组已完成课程要求的三项作品内容：论文 Word 文档、源文件文件夹和组员分工及其完成情况 Word 文档。源文件中包含 MySQL 5.5 兼容的数据库脚本和 PHP 源程序。")
    add_para(doc, "系统实现了零件信息的添加、修改、删除、查看和查询功能，数据库表设置了主码、数据类型和字段长度，插入了测试数据，满足信息管理系统基本功能要求。")
    doc.add_heading("三、协作说明", level=1)
    add_para(doc, "两名组员围绕“仓库零件管理系统”共同完成选题确认、字段设计和功能核对。组长负责整合最终材料，组员负责补充源程序和说明文件，最终统一形成可提交作品文件夹。")
    doc.save(OUT_DIR + r"\组员分工及其完成情况.docx")


if __name__ == "__main__":
    build_paper()
    build_assignment()
