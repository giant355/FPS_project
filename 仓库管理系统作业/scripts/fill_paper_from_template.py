from copy import deepcopy
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TEMPLATE = r"C:\Users\Lenovo\Downloads\《资料管理信息系统》论文.docx"
OUT = r"C:\Users\Lenovo\Desktop\新建文件夹\My project (11)\仓库管理系统_最终提交\论文-仓库管理系统的设计与实现.docx"


def set_run_font(run, name="宋体", size=12, bold=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_paragraph_text(paragraph, text, size=12, bold=None, align=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def remove_from_marker(doc, marker):
    body = doc._body._element
    start = None
    for child in list(body):
        text = "".join(node.text or "" for node in child.iter() if node.tag == qn("w:t"))
        if marker in text:
            start = child
            break
    if start is None:
        raise RuntimeError("未找到正文起始标记")
    seen = False
    for child in list(body):
        if child is start:
            seen = True
        if seen:
            body.remove(child)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = borders.find(qn("w:" + name))
        if border is None:
            border = OxmlElement("w:" + name)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_text(p, text, size=16 if level == 1 else 14, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
        if widths:
            table.rows[0].cells[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i != len(row) - 1 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
            if widths:
                cells[i].width = widths[i]
    set_table_borders(table)
    return table


def fill_cover(doc):
    # 保留模板封面格式，只替换组员信息。
    if len(doc.tables) >= 2:
        doc.tables[1].cell(0, 1).text = "魏皓荣(2024113060524)"
        doc.tables[1].cell(1, 1).text = "张珈宸(2023116000227)"
        doc.tables[1].cell(2, 1).text = "2026年7月"
        for table in doc.tables[:2]:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            set_run_font(run, size=12)


def build():
    doc = Document(TEMPLATE)
    fill_cover(doc)
    remove_from_marker(doc, "XXXX系统的设计与实现")

    title = doc.add_paragraph(style="CM20")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(title, "仓库管理系统的设计与实现", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "1前言", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_para(doc, "随着企业和学校实训室物资数量不断增加，仓库零件的登记、查询和维护工作也变得更加频繁。传统纸质登记方式容易出现记录不完整、查找速度慢、库存数量更新不及时等问题，影响仓库管理效率。")
    add_para(doc, "本系统以仓库零件信息为管理对象，围绕零件号、零件名、零件颜色、零件数量和购买时间等字段建立数据库，通过简单的信息管理页面完成零件信息维护。")
    add_heading(doc, "1.2 研究意义", 2)
    add_para(doc, "仓库零件管理系统能够把零散的零件信息集中保存到数据库中，使管理员可以快速查看库存情况，及时修改数量和基础信息，减少重复登记与人工查找带来的错误。")
    add_para(doc, "从课程实践角度看，本系统覆盖数据库创建、表结构设计、主码设置、测试数据插入以及 SQL 增删改查等知识点，能够体现数据库原理与应用课程的综合训练要求。")
    add_heading(doc, "1.3 研究内容", 2)
    add_para(doc, "本文主要完成仓库零件管理系统的需求分析、功能模块设计、数据库设计和功能实现说明。系统采用 MySQL 5.5 保存数据，使用 PHP 页面实现零件信息添加、修改、删除、查看和查询功能。")

    add_heading(doc, "2相关技术", 1)
    add_heading(doc, "2.1 MySQL 5.5", 2)
    add_para(doc, "MySQL 是常用的关系型数据库管理系统，MySQL 5.5 支持主键约束、字符集设置、日期类型和常见 SQL 操作。本系统使用 MySQL 5.5 创建 warehouse_db 数据库和 parts 数据表。")
    add_heading(doc, "2.2 PHP", 2)
    add_para(doc, "PHP 是服务器端脚本语言，适合开发中小型信息管理系统。本系统使用 PHP 的 mysqli 方式连接 MySQL，并通过 SQL 语句完成数据访问。")
    add_heading(doc, "2.3 HTML与CSS", 2)
    add_para(doc, "HTML 用于组织列表、表单和详情页面结构，CSS 用于设置页面布局、表格、按钮和输入框样式，使系统界面更清晰。")
    add_heading(doc, "2.4 CRUD功能", 2)
    add_para(doc, "CRUD 指创建、读取、更新和删除四类基本数据操作。本系统围绕零件信息完成添加、查看、修改和删除，满足信息管理系统的基本功能要求。")

    add_heading(doc, "3系统总体设计", 1)
    add_heading(doc, "3.1 需求分析", 2)
    add_para(doc, "系统面向仓库管理员使用，主要管理仓库中的零件基础信息。每条记录必须包含零件号、零件名、零件颜色、零件数量和购买时间，其中零件号作为唯一标识。")
    add_para(doc, "系统应实现以下功能：零件信息的添加、零件信息的修改、零件信息的删除、零件信息的查看。同时系统提供简单关键字查询，方便按零件号、零件名或颜色查找记录。")

    add_heading(doc, "3.2 功能模块设计", 2)
    add_para(doc, "从管理员使用角度分析，仓库零件管理系统主要由零件列表与查询模块、添加模块、修改模块、删除模块和查看模块组成。系统功能结构如图3-1所示。")
    add_caption(doc, "图 3-1 仓库零件管理系统功能结构图")
    add_para(doc, "1 零件信息添加：管理员录入零件号、零件名、零件颜色、零件数量和购买时间，系统校验后写入数据库。")
    add_para(doc, "2 零件信息维护：管理员可以根据零件号查看详情，对名称、颜色、数量、购买时间和备注进行修改，也可以删除不再使用的零件记录。")

    add_heading(doc, "3.3 数据库设计", 2)
    p = doc.add_paragraph()
    set_paragraph_text(p, "3.3.1 概念结构设计", size=12, bold=True)
    add_para(doc, "(1) 零件实体属性包括零件号、零件名、零件颜色、零件数量、购买时间、备注、创建时间和修改时间。零件实体属性图如图3-2所示。")
    add_caption(doc, "图 3-2 零件实体属性图")
    add_para(doc, "(2) 仓库管理员对零件实体执行添加、修改、删除和查看操作。系统核心对象单一，数据结构清晰，适合采用一张零件信息表完成管理。")
    add_caption(doc, "图 3-3 系统E-R图")

    p = doc.add_paragraph()
    set_paragraph_text(p, "3.3.2 逻辑结构设计", size=12, bold=True)
    add_para(doc, "本数据库采用 MYSQL 数据库。数据库名称为 warehouse_db，核心数据表为 parts。parts 表以 part_no 作为主键，保证零件号唯一。")
    add_para(doc, "(1) 零件信息表包括零件号、零件名、零件颜色、零件数量、购买时间、备注、创建时间、修改时间等字段，此表共设有8个字段，part_no为主键。如表3-1所示。")
    add_caption(doc, "表3-1 零件信息表")
    add_table(doc, ["字段名", "类型", "长度", "是否主键", "备注"], [
        ["part_no", "varchar", "20", "是", "零件号"],
        ["part_name", "varchar", "50", "否", "零件名"],
        ["part_color", "varchar", "20", "否", "零件颜色"],
        ["quantity", "int", "-", "否", "零件数量"],
        ["purchase_time", "date", "-", "否", "购买时间"],
        ["remark", "varchar", "200", "否", "备注"],
        ["created_at", "datetime", "-", "否", "创建时间"],
        ["updated_at", "datetime", "-", "否", "修改时间"],
    ], [Cm(3.0), Cm(2.5), Cm(1.8), Cm(2.2), Cm(5.0)])

    add_heading(doc, "4 系统的功能实现", 1)
    add_heading(doc, "4.1 数据库创建实现", 2)
    add_para(doc, "数据库脚本文件为 database/warehouse.sql。脚本首先创建 warehouse_db 数据库，然后创建 parts 表，并设置 part_no 为主码，最后插入若干测试数据，便于系统展示和查询。")
    add_heading(doc, "4.2 零件信息添加实现", 2)
    add_para(doc, "添加功能由 part_add.php 实现。管理员填写表单并提交后，程序检查必填字段和数量是否合法，再使用 INSERT INTO parts 语句把数据保存到数据库。")
    add_heading(doc, "4.3 零件信息修改实现", 2)
    add_para(doc, "修改功能由 part_edit.php 实现。系统根据零件号读取原数据并显示到表单，管理员修改后提交，程序通过 UPDATE parts SET ... WHERE part_no=... 更新指定记录。")
    add_heading(doc, "4.4 零件信息删除实现", 2)
    add_para(doc, "删除功能由 part_delete.php 实现。列表页面提供删除链接，并在浏览器中弹出确认提示。确认后系统根据零件号执行 DELETE FROM parts WHERE part_no=... 删除记录。")
    add_heading(doc, "4.5 查看与查询实现", 2)
    add_para(doc, "查看功能由 index.php 和 part_view.php 实现。index.php 显示零件列表，并支持按零件号、零件名、颜色进行模糊查询；part_view.php 显示单条零件记录的详细信息。")
    add_caption(doc, "表4-1 系统功能与主要文件对应表")
    add_table(doc, ["功能", "主要文件", "实现说明"], [
        ["查询/列表", "index.php", "使用 SELECT 语句读取零件记录并显示"],
        ["添加", "part_add.php", "使用 INSERT 语句新增零件记录"],
        ["修改", "part_edit.php", "使用 UPDATE 语句更新零件记录"],
        ["删除", "part_delete.php", "使用 DELETE 语句删除零件记录"],
        ["查看详情", "part_view.php", "根据零件号读取单条记录"],
    ], [Cm(3.0), Cm(3.8), Cm(8.7)])

    add_heading(doc, "5 总结", 1)
    add_para(doc, "本作品完成了仓库零件管理系统的设计与实现。系统围绕零件号、零件名、零件颜色、零件数量和购买时间等数据进行管理，实现了添加、修改、删除、查看和查询等基本功能。")
    add_para(doc, "通过本次综合作品，进一步掌握了 MySQL 5.5 数据库建库建表、字段类型选择、主键设置、测试数据插入和 SQL 增删改查语句的使用方法，也加深了对信息管理系统开发流程的理解。")

    doc.save(OUT)


if __name__ == "__main__":
    build()
