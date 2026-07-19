from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "paper_template_assets"
SOURCE_PATH = Path("C:/Users/Lenovo/Downloads/\u8bba\u6587.docx")
TEMPLATE_PATH = Path(
    "C:/Users/Lenovo/Downloads/\u300a\u8d44\u6599\u7ba1\u7406\u4fe1\u606f\u7cfb\u7edf\u300b\u8bba\u6587.docx"
)
OUT_PATH = OUT_DIR / "\u8d44\u6599\u7ba1\u7406\u4fe1\u606f\u7cfb\u7edf\u8bba\u6587-\u6309\u6a21\u677f\u586b\u5199.docx"


def ensure_dirs():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)


def get_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def set_run_font(run, east_asia="SimSun", ascii_font="Times New Roman", size=None):
    run.font.name = ascii_font
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)


def set_paragraph_format(paragraph, first_line=True):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.32)


def add_text(doc, text, style=None, align=None, first_line=True):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_format(paragraph, first_line=first_line)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="SimHei", ascii_font="Times New Roman", size=16 if level == 1 else 14)
    run.bold = True
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    return paragraph


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.text = ""
    shade_cell(cell, "F7F7F7")
    for raw_line in code.strip("\n").splitlines():
        paragraph = cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(raw_line)
        set_run_font(run, east_asia="SimSun", ascii_font="Consolas", size=9)
    if cell.paragraphs and not cell.paragraphs[0].text:
        cell._tc.remove(cell.paragraphs[0]._p)
    set_table_borders(table)
    return table


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "999999")


def add_table_from_rows(doc, rows, header=True):
    if not rows:
        return None
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 or len(value) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            set_run_font(run, size=10.5)
            if header and r == 0:
                run.bold = True
                shade_cell(cell, "EDEDED")
    doc.add_paragraph()
    return table


def table_rows(table):
    rows = []
    for row in table.rows:
        rows.append([" ".join(cell.text.split()) for cell in row.cells])
    return rows


def clear_body_from_title(doc):
    body = doc.element.body
    start = None
    for i, child in enumerate(list(body)):
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if "XXXX\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5b9e\u73b0" in text:
                start = i
                break
    if start is None:
        raise RuntimeError("Cannot locate template body title.")
    elements = list(body)
    for child in elements[start:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def update_cover_tables(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                key = "".join(cells[0].text.split())
                if key == "\u5b8c\u6210\u65e5\u671f":
                    cells[1].text = "2026\u5e746\u670830\u65e5"
                elif key == "\u7ec4\u957f":
                    cells[1].text = "\u8d56\u9e4f\u81e3(2024116050508)"
                elif key == "\u6210\u5458":
                    cells[1].text = "\u59ec\u9e4f(2024116050507)"


def draw_box(draw, xy, text, fill, outline="#4A5568", text_color="#111827", font=None, radius=12):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)
    x0, y0, x1, y1 = xy
    font = font or get_font(24, True)
    lines = text.split("\n")
    total_h = len(lines) * 30
    y = y0 + (y1 - y0 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x0 + (x1 - x0 - (bbox[2] - bbox[0])) / 2
        draw.text((x, y), line, fill=text_color, font=font)
        y += 30


def arrow(draw, start, end, color="#374151"):
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 > x1:
        pts = [(x2, y2), (x2 - 14, y2 - 9), (x2 - 14, y2 + 9)]
    elif x2 < x1:
        pts = [(x2, y2), (x2 + 14, y2 - 9), (x2 + 14, y2 + 9)]
    elif y2 > y1:
        pts = [(x2, y2), (x2 - 9, y2 - 14), (x2 + 9, y2 - 14)]
    else:
        pts = [(x2, y2), (x2 - 9, y2 + 14), (x2 + 9, y2 + 14)]
    draw.polygon(pts, fill=color)


def create_architecture_image():
    path = ASSET_DIR / "system_architecture.png"
    img = Image.new("RGB", (1150, 620), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = get_font(30, True)
    font = get_font(22, True)
    small = get_font(18)
    draw.text((360, 30), "\u8d44\u6599\u7ba1\u7406\u7cfb\u7edf\u6280\u672f\u67b6\u6784", fill="#111827", font=title_font)
    draw_box(draw, (70, 180, 270, 305), "\u6d4f\u89c8\u5668\nHTML/CSS/JS", "#E8F0FE", font=font)
    draw_box(draw, (370, 160, 610, 325), "Flask \u8def\u7531\u4e0e\n\u4e1a\u52a1\u6821\u9a8c", "#ECFDF5", font=font)
    draw_box(draw, (710, 160, 930, 325), "MySQLRepository\nPyMySQL", "#FFF7ED", font=font)
    draw_box(draw, (985, 180, 1110, 305), "MySQL\n5.5", "#FEE2E2", font=font)
    arrow(draw, (270, 242), (370, 242))
    arrow(draw, (610, 242), (710, 242))
    arrow(draw, (930, 242), (985, 242))
    draw_box(draw, (365, 420, 615, 520), "Jinja2 \u6a21\u677f\u6e32\u67d3\n\u8fd4\u56de\u9875\u9762", "#F3F4F6", font=font)
    arrow(draw, (490, 325), (490, 420))
    arrow(draw, (370, 470), (270, 305))
    draw.text((85, 350), "\u7528\u6237\u63d0\u4ea4\u67e5\u8be2\u3001\u6dfb\u52a0\u3001\u5220\u9664\u548c\u5206\u7ec4\u7ba1\u7406\u8bf7\u6c42", fill="#4B5563", font=small)
    img.save(path)
    return path


def create_er_image():
    path = ASSET_DIR / "er_diagram.png"
    img = Image.new("RGB", (1150, 620), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = get_font(30, True)
    font = get_font(22, True)
    small = get_font(18)
    draw.text((390, 30), "\u8d44\u6599\u7ba1\u7406\u7cfb\u7edf E-R \u56fe", fill="#111827", font=title_font)
    left_box = (110, 155, 430, 420)
    right_box = (720, 120, 1040, 470)
    draw.rounded_rectangle(left_box, radius=12, fill="#E8F0FE", outline="#4A5568", width=2)
    draw.rounded_rectangle(right_box, radius=12, fill="#ECFDF5", outline="#4A5568", width=2)
    draw.text((165, 185), "\u8d44\u6599\u5206\u7ec4 categories", fill="#111827", font=font)
    draw.line([(135, 230), (405, 230)], fill="#94A3B8", width=2)
    for i, line in enumerate(["id\uff08\u4e3b\u952e\uff09", "name\uff08\u552f\u4e00\uff09", "created_at"]):
        draw.text((160, 260 + i * 42), line, fill="#374151", font=small)
    draw.text((785, 150), "\u8d44\u6599 materials", fill="#111827", font=font)
    draw.line([(745, 195), (1015, 195)], fill="#94A3B8", width=2)
    for i, line in enumerate(["id\uff08\u4e3b\u952e\uff09", "title\uff08\u8d44\u6599\u540d\u79f0\uff09", "category_id\uff08\u5916\u952e\uff09", "quantity\uff08\u6570\u91cf\uff09", "content\uff08\u5185\u5bb9\uff09", "created_at / updated_at"]):
        draw.text((765, 220 + i * 38), line, fill="#374151", font=small)
    draw.line([(430, 285), (720, 285)], fill="#374151", width=4)
    draw.text((530, 245), "1", fill="#111827", font=font)
    draw.text((610, 245), "N", fill="#111827", font=font)
    draw.text((515, 312), "\u4e00\u4e2a\u5206\u7ec4\u5305\u542b\u591a\u6761\u8d44\u6599", fill="#4B5563", font=small)
    img.save(path)
    return path


def add_source_table(doc, source_doc, index, caption):
    add_caption(doc, caption)
    return add_table_from_rows(doc, table_rows(source_doc.tables[index]))


def build_document():
    ensure_dirs()
    source_doc = Document(SOURCE_PATH)
    doc = Document(TEMPLATE_PATH)
    update_cover_tables(doc)
    clear_body_from_title(doc)

    arch_img = create_architecture_image()
    er_img = create_er_image()

    title_style = "CM20" if "CM20" in [s.name for s in doc.styles] else None
    title = doc.add_paragraph(style=title_style)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("\u8d44\u6599\u7ba1\u7406\u4fe1\u606f\u7cfb\u7edf\u7684\u8bbe\u8ba1\u4e0e\u5b9e\u73b0")
    set_run_font(run, east_asia="SimHei", ascii_font="Times New Roman", size=18)
    run.bold = True

    add_heading(doc, "\u6458  \u8981", 1)
    add_text(
        doc,
        "\u968f\u7740\u7535\u5b50\u8d44\u6599\u6570\u91cf\u4e0d\u65ad\u589e\u52a0\uff0c\u4f7f\u7528\u6587\u4ef6\u5939\u548c\u7eb8\u8d28\u767b\u8bb0\u65b9\u5f0f\u5bb9\u6613\u51fa\u73b0\u5206\u7c7b\u6df7\u4e71\u3001\u6570\u91cf\u4e0d\u6e05\u548c\u67e5\u627e\u56f0\u96be\u7b49\u95ee\u9898\u3002\u672c\u6587\u8bbe\u8ba1\u5e76\u5b9e\u73b0\u4e00\u4e2a\u57fa\u4e8e Flask \u4e0e MySQL \u7684\u8d44\u6599\u7ba1\u7406\u4fe1\u606f\u7cfb\u7edf\u3002\u7cfb\u7edf\u56f4\u7ed5\u8d44\u6599\u79cd\u7c7b\u3001\u8d44\u6599\u6570\u91cf\u548c\u8d44\u6599\u5185\u5bb9\u4e09\u4e2a\u6838\u5fc3\u6570\u636e\u9879\uff0c\u63d0\u4f9b\u8d44\u6599\u6dfb\u52a0\u3001\u8d44\u6599\u5220\u9664\u3001\u8d44\u6599\u5206\u7ec4\u548c\u8d44\u6599\u67e5\u770b\u56db\u9879\u57fa\u672c\u529f\u80fd\uff0c\u5e76\u589e\u52a0\u641c\u7d22\u3001\u5206\u7ec4\u7b5b\u9009\u548c\u7edf\u8ba1\u603b\u89c8\u529f\u80fd\u3002\u7cfb\u7edf\u91c7\u7528\u6d4f\u89c8\u5668/\u670d\u52a1\u5668\u7ed3\u6784\uff0c\u8868\u73b0\u5c42\u4f7f\u7528 HTML\u3001CSS \u548c JavaScript\uff0c\u4e1a\u52a1\u5c42\u4f7f\u7528 Python Flask\uff0c\u6570\u636e\u5c42\u4f7f\u7528 PyMySQL \u8bbf\u95ee MySQL 5.5\u3002\u6570\u636e\u5e93\u7531\u8d44\u6599\u5206\u7ec4\u8868\u548c\u8d44\u6599\u8868\u6784\u6210\uff0c\u901a\u8fc7\u5916\u952e\u4fdd\u8bc1\u5206\u7c7b\u5173\u7cfb\u5b8c\u6574\uff0c\u901a\u8fc7\u540e\u7aef\u8f93\u5165\u6821\u9a8c\u4fdd\u8bc1\u8d44\u6599\u6570\u91cf\u4e3a\u6b63\u6574\u6570\u3002\u6d4b\u8bd5\u7ed3\u679c\u8868\u660e\uff0c\u7cfb\u7edf\u80fd\u591f\u6ee1\u8db3\u8bfe\u7a0b\u8bbe\u8ba1\u4e2d\u5bf9\u6570\u636e\u589e\u5220\u67e5\u3001\u5206\u7ec4\u7ba1\u7406\u548c\u6570\u636e\u5e93\u7ea6\u675f\u8bbe\u8ba1\u7684\u57fa\u672c\u8981\u6c42\u3002",
    )
    add_text(doc, "\u5173\u952e\u8bcd\uff1a\u8d44\u6599\u7ba1\u7406\uff1bFlask\uff1bMySQL\uff1b\u6570\u636e\u5e93\u8bbe\u8ba1\uff1bWeb\u7cfb\u7edf", first_line=False)

    add_heading(doc, "1\u524d\u8a00", 1)
    add_heading(doc, "1.1 \u7814\u7a76\u80cc\u666f", 2)
    add_text(
        doc,
        "\u5728\u8bfe\u7a0b\u5b66\u4e60\u3001\u5c0f\u7ec4\u9879\u76ee\u548c\u4e2a\u4eba\u77e5\u8bc6\u79ef\u7d2f\u8fc7\u7a0b\u4e2d\uff0c\u5b66\u751f\u4f1a\u4fdd\u5b58\u6559\u6750\u3001\u5b9e\u9a8c\u62a5\u544a\u3001\u9879\u76ee\u6587\u6863\u548c\u7535\u5b50\u4e66\u7b49\u591a\u79cd\u8d44\u6599\u3002\u5f53\u8d44\u6599\u6570\u91cf\u8f83\u5c11\u65f6\uff0c\u53ef\u4ee5\u4f9d\u9760\u6587\u4ef6\u5939\u540d\u79f0\u8fdb\u884c\u7ba1\u7406\uff1b\u5f53\u8d44\u6599\u9010\u6e10\u589e\u591a\u540e\uff0c\u5bb9\u6613\u51fa\u73b0\u540c\u7c7b\u8d44\u6599\u5206\u6563\u3001\u6570\u91cf\u7edf\u8ba1\u4e0d\u51c6\u786e\u3001\u5185\u5bb9\u8bf4\u660e\u7f3a\u5931\u7b49\u95ee\u9898\u3002\u56e0\u6b64\uff0c\u9700\u8981\u4e00\u4e2a\u7ed3\u6784\u7b80\u5355\u3001\u64cd\u4f5c\u76f4\u89c2\u7684\u8d44\u6599\u7ba1\u7406\u7cfb\u7edf\uff0c\u5c06\u8d44\u6599\u7684\u57fa\u672c\u4fe1\u606f\u7edf\u4e00\u4fdd\u5b58\u5728\u6570\u636e\u5e93\u4e2d\u3002",
    )
    add_heading(doc, "1.2 \u7814\u7a76\u610f\u4e49", 2)
    add_text(
        doc,
        "\u8d44\u6599\u7ba1\u7406\u4fe1\u606f\u7cfb\u7edf\u53ef\u4ee5\u628a\u5206\u6563\u7684\u5b66\u4e60\u548c\u9879\u76ee\u8d44\u6599\u6309\u7c7b\u522b\u7edf\u4e00\u7ec4\u7ec7\uff0c\u5e76\u901a\u8fc7\u6570\u636e\u5e93\u4fdd\u5b58\u8d44\u6599\u540d\u79f0\u3001\u8d44\u6599\u6570\u91cf\u548c\u8d44\u6599\u5185\u5bb9\u7b49\u4fe1\u606f\u3002\u8be5\u7cfb\u7edf\u4e0d\u4ec5\u80fd\u63d0\u9ad8\u67e5\u627e\u6548\u7387\uff0c\u4e5f\u80fd\u5c06\u6570\u636e\u5e93\u8bbe\u8ba1\u4e2d\u7684\u4e3b\u952e\u3001\u5916\u952e\u3001\u552f\u4e00\u7ea6\u675f\u548c\u5b8c\u6574\u6027\u6821\u9a8c\u7b49\u77e5\u8bc6\u8f6c\u5316\u4e3a\u53ef\u8fd0\u884c\u7684 Web \u5e94\u7528\uff0c\u5177\u6709\u8f83\u5f3a\u7684\u8bfe\u7a0b\u5b9e\u8df5\u610f\u4e49\u3002",
    )
    add_heading(doc, "1.3 \u7814\u7a76\u5185\u5bb9", 2)
    add_text(
        doc,
        "\u672c\u9879\u76ee\u76ee\u6807\u662f\u5b8c\u6210\u4e00\u4e2a\u80fd\u591f\u5b9e\u9645\u5c55\u793a\u6570\u636e\u5e93\u589e\u5220\u67e5\u529f\u80fd\u7684 Web \u7cfb\u7edf\u3002\u7cfb\u7edf\u56f4\u7ed5\u8d44\u6599\u79cd\u7c7b\u3001\u8d44\u6599\u6570\u91cf\u548c\u8d44\u6599\u5185\u5bb9\u5efa\u7acb\u6570\u636e\u6a21\u578b\uff0c\u5b9e\u73b0\u8d44\u6599\u6dfb\u52a0\u3001\u5220\u9664\u3001\u5206\u7ec4\u548c\u67e5\u770b\uff0c\u540c\u65f6\u4fdd\u8bc1\u8f93\u5165\u6570\u636e\u6709\u6548\u3001\u6570\u636e\u5e93\u5173\u7cfb\u5b8c\u6574\uff0c\u5e76\u63d0\u4f9b\u7b80\u6d01\u4e14\u9002\u914d\u7535\u8111\u4e0e\u624b\u673a\u7684\u7f51\u9875\u754c\u9762\u3002",
    )

    add_heading(doc, "2\u76f8\u5173\u6280\u672f", 1)
    add_heading(doc, "2.1 Flask", 2)
    add_text(doc, "Flask \u662f Python \u8bed\u8a00\u4e2d\u8f7b\u91cf\u7ea7 Web \u6846\u67b6\uff0c\u9002\u5408\u4e2d\u5c0f\u578b\u8bfe\u7a0b\u9879\u76ee\u5f00\u53d1\u3002\u672c\u7cfb\u7edf\u4f7f\u7528 Flask \u5b9a\u4e49\u8def\u7531\u3001\u63a5\u6536\u8868\u5355\u6570\u636e\u3001\u8fdb\u884c\u8f93\u5165\u6821\u9a8c\uff0c\u5e76\u8c03\u7528\u6570\u636e\u8bbf\u95ee\u5c42\u5b8c\u6210\u6570\u636e\u5e93\u8bfb\u5199\u3002")
    add_heading(doc, "2.2 B/S\u4e0eMVC\u6a21\u5f0f", 2)
    add_text(doc, "\u7cfb\u7edf\u91c7\u7528\u6d4f\u89c8\u5668/\u670d\u52a1\u5668\u7ed3\u6784\uff0c\u7528\u6237\u901a\u8fc7\u6d4f\u89c8\u5668\u8bbf\u95ee\u9875\u9762\uff0c\u540e\u7aef\u6839\u636e\u8def\u7531\u6267\u884c\u4e1a\u52a1\u5904\u7406\u3002\u5728\u5b9e\u73b0\u4e0a\uff0c\u7cfb\u7edf\u5c06\u8def\u7531\u63a7\u5236\u3001\u6570\u636e\u8bbf\u95ee\u3001HTML \u6a21\u677f\u548c\u9759\u6001\u6837\u5f0f\u5206\u79bb\uff0c\u4f53\u73b0\u4e86 MVC \u601d\u60f3\u4e2d\u804c\u8d23\u5206\u5c42\u7684\u57fa\u672c\u539f\u5219\u3002")
    add_heading(doc, "2.3 MySQL\u6570\u636e\u5e93", 2)
    add_text(doc, "\u672c\u7cfb\u7edf\u4f7f\u7528 MySQL 5.5 \u4f5c\u4e3a\u6570\u636e\u6301\u4e45\u5316\u5de5\u5177\u3002\u6570\u636e\u5e93\u4e2d\u5efa\u7acb\u8d44\u6599\u5206\u7ec4\u8868\u548c\u8d44\u6599\u8868\uff0c\u4f7f\u7528\u4e3b\u952e\u4fdd\u8bc1\u8bb0\u5f55\u552f\u4e00\uff0c\u4f7f\u7528\u5916\u952e\u7ef4\u62a4\u5206\u7ec4\u4e0e\u8d44\u6599\u4e4b\u95f4\u7684\u5f15\u7528\u5173\u7cfb\uff0c\u5e76\u901a\u8fc7\u7d22\u5f15\u63d0\u9ad8\u5206\u7ec4\u67e5\u8be2\u548c\u65f6\u95f4\u6392\u5e8f\u6548\u7387\u3002")
    add_heading(doc, "2.4 PyMySQL\u4e0eJinja2", 2)
    add_text(doc, "PyMySQL \u8d1f\u8d23\u5728 Python \u7a0b\u5e8f\u4e2d\u8fde\u63a5 MySQL\uff0c\u6267\u884c\u53c2\u6570\u5316 SQL \u548c\u4e8b\u52a1\u63a7\u5236\u3002Jinja2 \u8d1f\u8d23\u5c06\u540e\u7aef\u67e5\u8be2\u5230\u7684\u6570\u636e\u6e32\u67d3\u5230 HTML \u9875\u9762\u4e2d\uff0c\u4f7f\u7528\u6237\u80fd\u591f\u5728\u524d\u7aef\u67e5\u770b\u8d44\u6599\u5217\u8868\u3001\u8d44\u6599\u8be6\u60c5\u548c\u7edf\u8ba1\u4fe1\u606f\u3002")
    add_source_table(doc, source_doc, 1, "\u88682-1 \u7cfb\u7edf\u91c7\u7528\u6280\u672f\u4e0e\u4f5c\u7528")

    add_heading(doc, "3\u7cfb\u7edf\u603b\u4f53\u8bbe\u8ba1", 1)
    add_heading(doc, "3.1 \u9700\u6c42\u5206\u6790", 2)
    add_text(doc, "\u7cfb\u7edf\u9762\u5411\u9700\u8981\u6574\u7406\u5b66\u4e60\u8d44\u6599\u548c\u9879\u76ee\u8d44\u6599\u7684\u666e\u901a\u7528\u6237\u3002\u7528\u6237\u8fdb\u5165\u9996\u9875\u540e\u53ef\u4ee5\u67e5\u770b\u8d44\u6599\u7edf\u8ba1\u548c\u8d44\u6599\u5217\u8868\uff1b\u53ef\u901a\u8fc7\u5206\u7ec4\u6216\u5173\u952e\u8bcd\u7f29\u5c0f\u8303\u56f4\uff1b\u9700\u8981\u5f55\u5165\u65b0\u8d44\u6599\u65f6\u8fdb\u5165\u6dfb\u52a0\u9875\u9762\uff1b\u9700\u8981\u4e86\u89e3\u5b8c\u6574\u5185\u5bb9\u65f6\u8fdb\u5165\u8be6\u60c5\u9875\uff1b\u786e\u8ba4\u8d44\u6599\u65e0\u6548\u540e\u53ef\u6267\u884c\u5220\u9664\u3002")
    add_source_table(doc, source_doc, 2, "\u88683-1 \u8d44\u6599\u7ba1\u7406\u7cfb\u7edf\u529f\u80fd\u9700\u6c42")
    add_heading(doc, "3.2 \u529f\u80fd\u6a21\u5757\u8bbe\u8ba1", 2)
    add_text(doc, "\u7cfb\u7edf\u4ee5\u9996\u9875\u603b\u89c8\u3001\u8d44\u6599\u7ef4\u62a4\u3001\u5206\u7ec4\u7ba1\u7406\u548c\u641c\u7d22\u7b5b\u9009\u4e3a\u4e3b\u8981\u6a21\u5757\u3002\u5404\u6a21\u5757\u56f4\u7ed5\u8d44\u6599\u6570\u636e\u7684\u5f55\u5165\u3001\u5c55\u793a\u3001\u7edf\u8ba1\u548c\u5220\u9664\u5c55\u5f00\uff0c\u5176\u4e2d\u5199\u64cd\u4f5c\u7531\u540e\u7aef\u8fdb\u884c\u53c2\u6570\u6821\u9a8c\u548c\u4e8b\u52a1\u5904\u7406\u3002")
    add_source_table(doc, source_doc, 3, "\u88683-2 \u7cfb\u7edf\u529f\u80fd\u6a21\u5757")
    doc.add_picture(str(arch_img), width=Inches(5.9))
    add_caption(doc, "\u56fe3-1 \u8d44\u6599\u7ba1\u7406\u7cfb\u7edf\u6280\u672f\u67b6\u6784")
    add_heading(doc, "3.3 \u6570\u636e\u5e93\u8bbe\u8ba1", 2)
    add_text(doc, "\u6570\u636e\u5e93\u5305\u542b\u201c\u8d44\u6599\u5206\u7ec4\u201d\u548c\u201c\u8d44\u6599\u201d\u4e24\u4e2a\u5b9e\u4f53\u3002\u4e00\u4e2a\u5206\u7ec4\u53ef\u4ee5\u5305\u542b\u591a\u6761\u8d44\u6599\uff0c\u4e00\u6761\u8d44\u6599\u5fc5\u987b\u4e14\u53ea\u80fd\u5c5e\u4e8e\u4e00\u4e2a\u5206\u7ec4\uff0c\u56e0\u6b64\u4e8c\u8005\u662f 1:N \u5173\u7cfb\u3002")
    doc.add_picture(str(er_img), width=Inches(5.9))
    add_caption(doc, "\u56fe3-2 \u8d44\u6599\u7ba1\u7406\u7cfb\u7edfE-R\u56fe")
    add_text(doc, "\u5173\u7cfb\u6a21\u5f0f\u4e3a\uff1a\u8d44\u6599\u5206\u7ec4\uff08id\uff0cname\uff0ccreated_at\uff09\uff1b\u8d44\u6599\uff08id\uff0ctitle\uff0ccategory_id\uff0cquantity\uff0ccontent\uff0ccreated_at\uff0cupdated_at\uff09\u3002\u5176\u4e2d category_id \u662f\u5916\u952e\uff0c\u5f15\u7528\u8d44\u6599\u5206\u7ec4\u8868\u7684 id\u3002")
    add_source_table(doc, source_doc, 4, "\u88683-3 \u8d44\u6599\u7ba1\u7406\u7cfb\u7edf\u6570\u636e\u8868\u8bbe\u8ba1")
    add_text(doc, "\u7cfb\u7edf\u540c\u65f6\u4f7f\u7528\u5b9e\u4f53\u5b8c\u6574\u6027\u3001\u53c2\u7167\u5b8c\u6574\u6027\u548c\u7528\u6237\u5b9a\u4e49\u5b8c\u6574\u6027\u3002\u4e24\u5f20\u8868\u5747\u8bbe\u7f6e\u81ea\u589e\u4e3b\u952e\uff1bmaterials.category_id \u901a\u8fc7\u5916\u952e\u5f15\u7528 categories.id\uff1b\u5206\u7ec4\u540d\u79f0\u8bbe\u7f6e\u552f\u4e00\u7d22\u5f15\uff1bquantity \u4f7f\u7528\u540e\u7aef\u8f93\u5165\u6821\u9a8c\u9650\u5236\u4e3a\u6b63\u6570\u3002")
    add_caption(doc, "\u88683-4 \u6570\u636e\u5e93\u521d\u59cb\u5316\u6838\u5fc3SQL")
    add_code_block(
        doc,
        """CREATE DATABASE material_manager DEFAULT CHARACTER SET utf8;
USE material_manager;

CREATE TABLE categories (
    id INT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    name VARCHAR(50) NOT NULL UNIQUE,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE materials (
    id INT UNSIGNED AUTO_INCREMENT PRIMARY KEY,
    title VARCHAR(120) NOT NULL,
    category_id INT UNSIGNED NOT NULL,
    quantity INT UNSIGNED NOT NULL,
    content TEXT NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
    FOREIGN KEY (category_id) REFERENCES categories(id)
);""",
    )

    add_heading(doc, "4 \u7cfb\u7edf\u7684\u529f\u80fd\u5b9e\u73b0", 1)
    add_heading(doc, "4.1 \u9996\u9875\u4e0e\u67e5\u770b\u6a21\u5757\u5b9e\u73b0", 2)
    add_text(doc, "\u9996\u9875\u8c03\u7528 list_materials\u3001list_categories \u548c get_stats \u4e09\u4e2a\u65b9\u6cd5\uff0c\u5206\u522b\u53d6\u5f97\u8d44\u6599\u5217\u8868\u3001\u5206\u7ec4\u7edf\u8ba1\u548c\u603b\u89c8\u6570\u636e\u3002\u7528\u6237\u53ef\u8f93\u5165\u5173\u952e\u8bcd\u641c\u7d22\u8d44\u6599\u540d\u79f0\u6216\u5185\u5bb9\uff0c\u4e5f\u53ef\u9009\u62e9\u5206\u7ec4\u8fdb\u884c\u7b5b\u9009\u3002\u8be6\u60c5\u9875\u6839\u636e\u8d44\u6599\u7f16\u53f7\u67e5\u8be2\u5355\u6761\u8bb0\u5f55\uff0c\u627e\u4e0d\u5230\u65f6\u8fd4\u56de 404 \u9875\u9762\u3002")
    add_code_block(
        doc,
        """@app.get("/")
def index():
    category_id = request.args.get("category", type=int)
    keyword = request.args.get("keyword", "").strip()
    return render_template(
        "index.html",
        materials=repo().list_materials(category_id, keyword),
        categories=repo().list_categories(),
        stats=repo().get_stats()
    )""",
    )
    add_heading(doc, "4.2 \u8d44\u6599\u6dfb\u52a0\u6a21\u5757\u5b9e\u73b0", 2)
    add_text(doc, "\u6dfb\u52a0\u8868\u5355\u5305\u542b\u8d44\u6599\u540d\u79f0\u3001\u8d44\u6599\u79cd\u7c7b\u3001\u8d44\u6599\u6570\u91cf\u548c\u8d44\u6599\u5185\u5bb9\u3002\u540e\u7aef\u518d\u6b21\u6821\u9a8c\u6240\u6709\u5b57\u6bb5\uff0c\u6570\u91cf\u5fc5\u987b\u4e3a\u5927\u4e8e 0 \u7684\u6574\u6570\u3002\u6821\u9a8c\u901a\u8fc7\u540e\u6267\u884c INSERT\uff0c\u5e76\u8df3\u8f6c\u5230\u65b0\u8d44\u6599\u8be6\u60c5\u9875\uff1b\u6821\u9a8c\u5931\u8d25\u5219\u4fdd\u7559\u7528\u6237\u8f93\u5165\u5e76\u63d0\u793a\u9519\u8bef\u3002")
    add_code_block(doc, "INSERT INTO materials (title, category_id, quantity, content)\nVALUES (%s, %s, %s, %s);")
    add_heading(doc, "4.3 \u8d44\u6599\u5220\u9664\u6a21\u5757\u5b9e\u73b0", 2)
    add_text(doc, "\u5220\u9664\u64cd\u4f5c\u4f7f\u7528 POST \u8bf7\u6c42\uff0c\u9875\u9762\u901a\u8fc7 JavaScript \u663e\u793a\u4e8c\u6b21\u786e\u8ba4\u3002\u6570\u636e\u5e93\u6839\u636e\u4e3b\u952e\u6267\u884c DELETE\uff0c\u5e76\u6839\u636e\u53d7\u5f71\u54cd\u884c\u6570\u53cd\u9988\u201c\u5220\u9664\u6210\u529f\u201d\u6216\u201c\u8bb0\u5f55\u4e0d\u5b58\u5728\u201d\u3002\u5220\u9664\u5c5e\u4e8e\u5199\u64cd\u4f5c\uff0c\u6210\u529f\u540e\u63d0\u4ea4\u4e8b\u52a1\uff0c\u5f02\u5e38\u65f6\u56de\u6eda\u3002")
    add_heading(doc, "4.4 \u8d44\u6599\u5206\u7ec4\u4e0e\u7b5b\u9009\u6a21\u5757\u5b9e\u73b0", 2)
    add_text(doc, "\u5206\u7ec4\u8868\u4f7f\u7528\u552f\u4e00\u7d22\u5f15\u9632\u6b62\u91cd\u540d\u3002\u9996\u9875\u5de6\u4fa7\u5c55\u793a\u6bcf\u4e2a\u5206\u7ec4\u7684\u8d44\u6599\u6761\u6570\u548c\u8d44\u6599\u603b\u6570\u91cf\u3002\u5220\u9664\u5206\u7ec4\u524d\u5148\u68c0\u67e5\u662f\u5426\u4ecd\u6709\u5173\u8054\u8d44\u6599\uff1b\u975e\u7a7a\u5206\u7ec4\u4e0d\u5141\u8bb8\u5220\u9664\uff0c\u4ece\u5e94\u7528\u5c42\u548c\u5916\u952e\u7ea6\u675f\u4e24\u65b9\u9762\u4fdd\u8bc1\u6570\u636e\u5b8c\u6574\u3002")
    add_heading(doc, "4.5 \u7cfb\u7edf\u6d4b\u8bd5", 2)
    add_text(doc, "\u9879\u76ee\u5728 tests/test_app.py \u4e2d\u63d0\u4f9b\u57fa\u4e8e unittest \u7684\u8def\u7531\u6d4b\u8bd5\u3002\u6d4b\u8bd5\u4f7f\u7528\u5185\u5b58\u4e2d\u7684 FakeRepository \u66ff\u4ee3\u771f\u5b9e MySQL \u8fde\u63a5\uff0c\u7528\u4e8e\u9a8c\u8bc1\u9875\u9762\u8def\u7531\u3001\u8f93\u5165\u6821\u9a8c\u548c\u4e1a\u52a1\u6d41\u7a0b\u3002MySQL \u90e8\u5206\u901a\u8fc7\u68c0\u67e5\u5efa\u8868\u811a\u672c\u3001\u5916\u952e\u3001\u552f\u4e00\u7ea6\u675f\u3001\u6570\u91cf\u540e\u7aef\u8f93\u5165\u6821\u9a8c\u548c\u53c2\u6570\u5316 SQL \u8fdb\u884c\u9759\u6001\u9a8c\u8bc1\u3002")
    add_source_table(doc, source_doc, 5, "\u88684-1 \u7cfb\u7edf\u529f\u80fd\u6d4b\u8bd5\u7528\u4f8b")
    add_text(doc, "\u8fd0\u884c\u73af\u5883\u65b9\u9762\uff0c\u7cfb\u7edf\u9700\u8981 Python 3.10 \u6216\u4ee5\u4e0a\u7248\u672c\u3001MySQL 5.5 \u548c Flask\u3001PyMySQL \u7b49 Python \u4f9d\u8d56\u3002\u9996\u6b21\u8fd0\u884c\u524d\u9700\u5148\u5bfc\u5165 material_manager.sql\uff0c\u518d\u914d\u7f6e\u6570\u636e\u5e93\u7528\u6237\u540d\u548c\u5bc6\u7801\uff0c\u6700\u540e\u542f\u52a8 Flask \u5e94\u7528\u5e76\u901a\u8fc7\u6d4f\u89c8\u5668\u8bbf\u95ee\u3002")
    add_source_table(doc, source_doc, 6, "\u88684-2 \u7cfb\u7edf\u90e8\u7f72\u6b65\u9aa4")

    add_heading(doc, "5 \u7ed3\u8bba", 1)
    add_text(doc, "\u672c\u9879\u76ee\u5b8c\u6210\u4e86\u4e00\u4e2a\u7ed3\u6784\u6e05\u695a\u7684\u8d44\u6599\u7ba1\u7406\u4fe1\u606f\u7cfb\u7edf\uff0c\u5b9e\u73b0\u8d44\u6599\u6dfb\u52a0\u3001\u8d44\u6599\u5220\u9664\u3001\u8d44\u6599\u5206\u7ec4\u548c\u8d44\u6599\u67e5\u770b\u56db\u9879\u57fa\u672c\u529f\u80fd\uff0c\u5e76\u8986\u76d6\u8d44\u6599\u79cd\u7c7b\u3001\u8d44\u6599\u6570\u91cf\u548c\u8d44\u6599\u5185\u5bb9\u4e09\u4e2a\u6838\u5fc3\u6570\u636e\u9879\u3002\u7cfb\u7edf\u91c7\u7528 Flask \u4e0e MySQL \u5206\u5c42\u5b9e\u73b0\uff0c\u6570\u636e\u5e93\u4f7f\u7528\u4e3b\u952e\u3001\u5916\u952e\u3001\u552f\u4e00\u7ea6\u675f\u548c\u540e\u7aef\u8f93\u5165\u6821\u9a8c\u4fdd\u8bc1\u6570\u636e\u8d28\u91cf\uff0c\u7a0b\u5e8f\u4f7f\u7528\u53c2\u6570\u5316 SQL \u548c\u4e8b\u52a1\u63d0\u9ad8\u5b89\u5168\u6027\u4e0e\u53ef\u9760\u6027\u3002")
    add_text(doc, "\u53d7\u8bfe\u7a0b\u89c4\u6a21\u9650\u5236\uff0c\u5f53\u524d\u7cfb\u7edf\u6ca1\u6709\u5b9e\u73b0\u7528\u6237\u767b\u5f55\u3001\u8d44\u6599\u9644\u4ef6\u4e0a\u4f20\u548c\u4fee\u6539\u529f\u80fd\u3002\u540e\u7eed\u53ef\u4ee5\u589e\u52a0\u8d26\u53f7\u6743\u9650\u3001\u6587\u4ef6\u5b58\u50a8\u3001\u8d44\u6599\u7f16\u8f91\u3001\u5206\u9875\u67e5\u8be2\u548c\u6570\u636e\u5bfc\u51fa\uff0c\u5e76\u5728\u771f\u5b9e MySQL \u670d\u52a1\u4e0a\u8865\u5145\u6027\u80fd\u6d4b\u8bd5\u3002")

    add_heading(doc, "\u53c2\u8003\u6587\u732e", 1)
    for item in [
        "[1] \u738b\u73ca, \u8428\u5e08\u714a. \u6570\u636e\u5e93\u7cfb\u7edf\u6982\u8bba\uff08\u7b2c5\u7248\uff09. \u9ad8\u7b49\u6559\u80b2\u51fa\u7248\u793e, 2014.",
        "[2] Flask Documentation. Pallets Projects.",
        "[3] MySQL 5.5 Reference Manual. Oracle.",
        "[4] Python 3 Documentation. Python Software Foundation.",
    ]:
        add_text(doc, item, first_line=False)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
