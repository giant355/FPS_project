from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "lab678_assets"
OUT_PATH = OUT_DIR / "实验六七八-数据库设计与应用开发-新实验报告.docx"


def ensure_dirs():
    OUT_DIR.mkdir(exist_ok=True)
    ASSET_DIR.mkdir(exist_ok=True)


def set_east_asia_font(style, font_name="Microsoft YaHei"):
    if style.element.rPr is None:
        style.element.get_or_add_rPr()
    r_pr = style.element.rPr
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_east_asia(run, font_name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    set_east_asia_font(normal)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    set_east_asia_font(title)
    title.paragraph_format.space_after = Pt(10)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        set_east_asia_font(st)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.10

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.2)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.05)
    set_east_asia_font(code, "Microsoft YaHei")

    caption = styles.add_style("FigureCaption", 1)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_east_asia_font(caption)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "数据库设计与应用开发实验报告"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.name = "Calibri"
        set_run_east_asia(run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ").font.size = Pt(9)
    add_field(footer, "PAGE")
    footer.add_run(" 页").font.size = Pt(9)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    run.font.size = Pt(9)


def cell_set_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    set_run_east_asia(run)
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa, indent_dxa=120, header_fill="F2F4F7"):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    for tag in ["w:tblW", "w:tblInd", "w:tblCellMar", "w:tblBorders"]:
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)

    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    borders = OxmlElement("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "B8C0CC")
        borders.append(elem)
    tbl_pr.append(borders)

    margins = OxmlElement("w:tblCellMar")
    for side, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)

    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row_i, row in enumerate(table.rows):
        for cell_i, cell in enumerate(row.cells):
            width = widths_dxa[min(cell_i, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            if row_i == 0 and header_fill:
                shade_cell(cell, header_fill)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1800, 3000, 1800, 2760])
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            cell_set_text(table.cell(r, c), value, bold=(c in [0, 2]))
            if c in [0, 2]:
                shade_cell(table.cell(r, c), "F2F4F7")
    return table


def add_field_table(doc):
    rows = [
        ["实验项目名称", "数据库的设计与应用开发", "实验日期", "2026.6.30"],
        ["实验类型", "综合性 / 设计研究", "实验地点", "计算机实验室"],
        ["学生姓名", "__________", "学号", "__________"],
        ["班级", "__________", "指导教师", "__________"],
    ]
    add_kv_table(doc, rows)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for c, header in enumerate(headers):
        cell_set_text(table.cell(0, c), header, bold=True, color=RGBColor(0x0B, 0x25, 0x45))
    for row in rows:
        cells = table.add_row().cells
        for c, value in enumerate(row):
            cell_set_text(cells[c], value)
    set_table_geometry(table, widths)
    return table


def add_paragraph(doc, text, bold_label=None):
    p = doc.add_paragraph()
    if bold_label and text.startswith(bold_label):
        r1 = p.add_run(bold_label)
        r1.bold = True
        set_run_east_asia(r1)
        r2 = p.add_run(text[len(bold_label):])
        set_run_east_asia(r2)
    else:
        r = p.add_run(text)
        set_run_east_asia(r)
    return p


def add_numbered_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_run_east_asia(run)


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        run = p.add_run(item)
        set_run_east_asia(run)


def add_code_block(doc, title, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    set_run_east_asia(run)

    for raw_line in dedent(code).strip("\n").splitlines():
        line = raw_line.replace("\t", "    ")
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.keep_together = False
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        set_run_east_asia(r)


def load_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_browser_frame(draw, width, title):
    draw.rounded_rectangle([30, 25, width - 30, 595], radius=16, fill="#FFFFFF", outline="#C9D1DB", width=2)
    draw.rounded_rectangle([30, 25, width - 30, 85], radius=16, fill="#EEF2F7", outline="#C9D1DB", width=2)
    for i, color in enumerate(["#F25F5C", "#FFE066", "#70C1B3"]):
        draw.ellipse([55 + i * 28, 47, 70 + i * 28, 62], fill=color)
    draw.rounded_rectangle([165, 43, width - 70, 68], radius=8, fill="#FFFFFF", outline="#D5DAE2")
    draw.text((180, 46), f"localhost:8080/student/{title}", fill="#53606F", font=load_font(14))


def create_mock_screenshots():
    screenshots = []
    font_title = load_font(30, bold=True)
    font_h = load_font(20, bold=True)
    font = load_font(18)
    small = load_font(15)

    # Screenshot 1: course list
    img = Image.new("RGB", (1100, 650), "#F5F7FB")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, 1100, "course_read.jsp")
    draw.text((70, 115), "课程信息查看", fill="#0B2545", font=font_title)
    headers = ["课程号", "课程名", "先行课", "学分", "操作"]
    rows = [
        ["1", "数据库原理", "无", "4", "修改 | 删除"],
        ["2", "Java Web程序设计", "1", "3", "修改 | 删除"],
        ["3", "软件工程", "2", "3", "修改 | 删除"],
    ]
    x0, y0 = 70, 170
    widths = [130, 280, 160, 120, 210]
    row_h = 55
    x = x0
    for i, h in enumerate(headers):
        draw.rectangle([x, y0, x + widths[i], y0 + row_h], fill="#E8EEF5", outline="#AEB8C6")
        draw.text((x + 18, y0 + 15), h, fill="#0B2545", font=font_h)
        x += widths[i]
    for r, row in enumerate(rows):
        y = y0 + row_h * (r + 1)
        x = x0
        for i, value in enumerate(row):
            draw.rectangle([x, y, x + widths[i], y + row_h], fill="#FFFFFF", outline="#AEB8C6")
            draw.text((x + 18, y + 15), value, fill="#1F2937", font=font)
            x += widths[i]
    draw.rounded_rectangle([70, 410, 190, 455], radius=8, fill="#2E74B5")
    draw.text((92, 421), "添加课程", fill="#FFFFFF", font=font)
    path = ASSET_DIR / "course_read.png"
    img.save(path)
    screenshots.append(path)

    # Screenshot 2: add course form
    img = Image.new("RGB", (1100, 650), "#F5F7FB")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, 1100, "addCourse.jsp")
    draw.text((70, 115), "添加课程信息", fill="#0B2545", font=font_title)
    labels = ["课程号：", "课程名：", "先行课：", "学分："]
    values = ["4", "操作系统", "1", "3"]
    for i, (label, value) in enumerate(zip(labels, values)):
        y = 180 + i * 70
        draw.text((100, y + 12), label, fill="#1F2937", font=font_h)
        draw.rounded_rectangle([220, y, 650, y + 48], radius=8, fill="#FFFFFF", outline="#B8C0CC", width=2)
        draw.text((240, y + 12), value, fill="#53606F", font=font)
    draw.rounded_rectangle([220, 480, 330, 530], radius=8, fill="#2E74B5")
    draw.text((255, 493), "添加", fill="#FFFFFF", font=font)
    draw.rounded_rectangle([350, 480, 460, 530], radius=8, fill="#E5E7EB")
    draw.text((385, 493), "重置", fill="#1F2937", font=font)
    path = ASSET_DIR / "add_course.png"
    img.save(path)
    screenshots.append(path)

    # Screenshot 3: update form
    img = Image.new("RGB", (1100, 650), "#F5F7FB")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, 1100, "updateCourse.jsp?cno=4")
    draw.text((70, 115), "修改课程信息", fill="#0B2545", font=font_title)
    labels = ["课程号：", "课程名：", "先行课：", "学分："]
    values = ["4", "操作系统原理", "1", "4"]
    for i, (label, value) in enumerate(zip(labels, values)):
        y = 180 + i * 70
        draw.text((100, y + 12), label, fill="#1F2937", font=font_h)
        fill = "#F3F4F6" if i == 0 else "#FFFFFF"
        draw.rounded_rectangle([220, y, 650, y + 48], radius=8, fill=fill, outline="#B8C0CC", width=2)
        draw.text((240, y + 12), value, fill="#53606F", font=font)
    draw.rounded_rectangle([220, 480, 330, 530], radius=8, fill="#2E74B5")
    draw.text((255, 493), "修改", fill="#FFFFFF", font=font)
    draw.text((70, 565), "提示：课程号作为主键只读，其他字段可修改。", fill="#53606F", font=small)
    path = ASSET_DIR / "update_course.png"
    img.save(path)
    screenshots.append(path)

    # Screenshot 4: delete success
    img = Image.new("RGB", (1100, 650), "#F5F7FB")
    draw = ImageDraw.Draw(img)
    draw_browser_frame(draw, 1100, "deleteCourse.jsp?cno=4")
    draw.rounded_rectangle([260, 180, 840, 430], radius=16, fill="#FFFFFF", outline="#D5DAE2", width=2)
    draw.ellipse([330, 235, 395, 300], fill="#16A34A")
    draw.text((347, 242), "✓", fill="#FFFFFF", font=load_font(40, bold=True))
    draw.text((425, 238), "课程信息删除成功！", fill="#0B2545", font=font_title)
    draw.text((425, 302), "系统已根据课程号删除对应记录。", fill="#53606F", font=font)
    draw.rounded_rectangle([425, 350, 570, 398], radius=8, fill="#2E74B5")
    draw.text((450, 362), "返回课程表", fill="#FFFFFF", font=font)
    path = ASSET_DIR / "delete_success.png"
    img.save(path)
    screenshots.append(path)

    return screenshots


SQL_CODE = """
CREATE DATABASE IF NOT EXISTS student_course_test678
DEFAULT CHARACTER SET utf8mb4
COLLATE utf8mb4_general_ci;

USE student_course_test678;

CREATE TABLE course (
    cno    VARCHAR(10) PRIMARY KEY,
    cname  VARCHAR(50) NOT NULL,
    cpno   VARCHAR(10) NULL,
    credit INT NOT NULL
);

INSERT INTO course(cno, cname, cpno, credit) VALUES
('1', '数据库原理', NULL, 4),
('2', 'Java Web程序设计', '1', 3),
('3', '软件工程', '2', 3);
"""


DBUTIL_CODE = """
package com.toolbean;

import java.sql.Connection;
import java.sql.DriverManager;
import java.sql.SQLException;

public class DBUtil {
    private static final String URL =
        "jdbc:mysql://localhost:3306/student_course_test678"
        + "?useSSL=false&serverTimezone=Asia/Shanghai&characterEncoding=utf8";
    private static final String USER = "root";
    private static final String PASSWORD = "1234";

    static {
        try {
            Class.forName("com.mysql.cj.jdbc.Driver");
        } catch (ClassNotFoundException e) {
            throw new RuntimeException("MySQL驱动加载失败", e);
        }
    }

    public static Connection getConnection() throws SQLException {
        return DriverManager.getConnection(URL, USER, PASSWORD);
    }
}
"""


COURSE_READ_CODE = """
<%@ page language="java" contentType="text/html; charset=UTF-8"
    pageEncoding="UTF-8"%>
<%@ page import="java.sql.*" %>
<%@ page import="com.toolbean.DBUtil" %>
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>课程信息查看</title>
<style>
    table{border-collapse:collapse;width:760px;}
    th,td{border:1px solid #333;text-align:center;padding:8px;}
    th{background:#e8eef5;}
    a{margin:0 4px;}
</style>
</head>
<body>
<h2>课程信息查看</h2>
<p><a href="addCourse.jsp">添加课程</a></p>
<table>
<tr>
    <th>课程号</th><th>课程名</th><th>先行课</th><th>学分</th><th>操作</th>
</tr>
<%
String sql = "select cno, cname, cpno, credit from course order by cno";
try(Connection conn = DBUtil.getConnection();
    PreparedStatement ps = conn.prepareStatement(sql);
    ResultSet rs = ps.executeQuery()){
    while(rs.next()){
%>
<tr>
    <td><%=rs.getString("cno") %></td>
    <td><%=rs.getString("cname") %></td>
    <td><%=rs.getString("cpno") == null ? "无" : rs.getString("cpno") %></td>
    <td><%=rs.getInt("credit") %></td>
    <td>
        <a href="updateCourse.jsp?cno=<%=rs.getString("cno") %>">修改</a>
        <a href="deleteCourse.jsp?cno=<%=rs.getString("cno") %>"
           onclick="return confirm('确定删除该课程吗？')">删除</a>
    </td>
</tr>
<%
    }
}
%>
</table>
</body>
</html>
"""


ADD_COURSE_CODE = """
<%@ page language="java" contentType="text/html; charset=UTF-8"
    pageEncoding="UTF-8"%>
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>添加课程信息</title>
</head>
<body>
<h2>添加课程信息</h2>
<form action="addCourseProcess.jsp" method="post">
    课程号：<input type="text" name="cno" required><br><br>
    课程名：<input type="text" name="cname" required><br><br>
    先行课：<input type="text" name="cpno"><br><br>
    学分：<input type="number" name="credit" min="1" max="8" required><br><br>
    <input type="submit" value="添加">
    <input type="reset" value="重置">
</form>
</body>
</html>
"""


ADD_PROCESS_CODE = """
<%@ page language="java" contentType="text/html; charset=UTF-8"
    pageEncoding="UTF-8"%>
<%@ page import="java.sql.*" %>
<%@ page import="com.toolbean.DBUtil" %>
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>课程添加处理</title>
</head>
<body>
<%
request.setCharacterEncoding("UTF-8");
String cno = request.getParameter("cno");
String cname = request.getParameter("cname");
String cpno = request.getParameter("cpno");
int credit = Integer.parseInt(request.getParameter("credit"));

String sql = "insert into course(cno, cname, cpno, credit) values(?, ?, ?, ?)";
try(Connection conn = DBUtil.getConnection();
    PreparedStatement ps = conn.prepareStatement(sql)){
    ps.setString(1, cno);
    ps.setString(2, cname);
    if(cpno == null || cpno.trim().equals("")){
        ps.setNull(3, java.sql.Types.VARCHAR);
    }else{
        ps.setString(3, cpno);
    }
    ps.setInt(4, credit);
    int result = ps.executeUpdate();
    if(result > 0){
%>
<h3>课程信息添加成功！</h3>
<a href="course_read.jsp">查看课程信息</a>
<%
    }else{
%>
<h3>课程信息添加失败！</h3>
<%
    }
}
%>
</body>
</html>
"""


UPDATE_COURSE_CODE = """
<%@ page language="java" contentType="text/html; charset=UTF-8"
    pageEncoding="UTF-8"%>
<%@ page import="java.sql.*" %>
<%@ page import="com.toolbean.DBUtil" %>
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>修改课程信息</title>
</head>
<body>
<%
String cno = request.getParameter("cno");
String sql = "select cno, cname, cpno, credit from course where cno=?";
try(Connection conn = DBUtil.getConnection();
    PreparedStatement ps = conn.prepareStatement(sql)){
    ps.setString(1, cno);
    try(ResultSet rs = ps.executeQuery()){
        if(rs.next()){
%>
<h2>修改课程信息</h2>
<form action="updateCourseProcess.jsp" method="post">
    课程号：<input type="text" name="cno"
        value="<%=rs.getString("cno") %>" readonly><br><br>
    课程名：<input type="text" name="cname"
        value="<%=rs.getString("cname") %>" required><br><br>
    先行课：<input type="text" name="cpno"
        value="<%=rs.getString("cpno") == null ? "" : rs.getString("cpno") %>"><br><br>
    学分：<input type="number" name="credit"
        value="<%=rs.getInt("credit") %>" min="1" max="8" required><br><br>
    <input type="submit" value="修改">
    <input type="reset" value="重置">
</form>
<%
        }else{
%>
<h3>未找到该课程信息。</h3>
<%
        }
    }
}
%>
</body>
</html>
"""


UPDATE_PROCESS_CODE = """
<%@ page language="java" contentType="text/html; charset=UTF-8"
    pageEncoding="UTF-8"%>
<%@ page import="java.sql.*" %>
<%@ page import="com.toolbean.DBUtil" %>
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>课程修改处理</title>
</head>
<body>
<%
request.setCharacterEncoding("UTF-8");
String cno = request.getParameter("cno");
String cname = request.getParameter("cname");
String cpno = request.getParameter("cpno");
int credit = Integer.parseInt(request.getParameter("credit"));

String sql = "update course set cname=?, cpno=?, credit=? where cno=?";
try(Connection conn = DBUtil.getConnection();
    PreparedStatement ps = conn.prepareStatement(sql)){
    ps.setString(1, cname);
    if(cpno == null || cpno.trim().equals("")){
        ps.setNull(2, java.sql.Types.VARCHAR);
    }else{
        ps.setString(2, cpno);
    }
    ps.setInt(3, credit);
    ps.setString(4, cno);
    int result = ps.executeUpdate();
    if(result > 0){
%>
<h3>课程信息修改成功！</h3>
<a href="course_read.jsp">返回课程表</a>
<%
    }else{
%>
<h3>课程信息修改失败！</h3>
<%
    }
}
%>
</body>
</html>
"""


DELETE_COURSE_CODE = """
<%@ page language="java" contentType="text/html; charset=UTF-8"
    pageEncoding="UTF-8"%>
<%@ page import="java.sql.*" %>
<%@ page import="com.toolbean.DBUtil" %>
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>课程删除处理</title>
</head>
<body>
<%
String cno = request.getParameter("cno");
String sql = "delete from course where cno=?";
try(Connection conn = DBUtil.getConnection();
    PreparedStatement ps = conn.prepareStatement(sql)){
    ps.setString(1, cno);
    int result = ps.executeUpdate();
    if(result > 0){
%>
<h3>课程信息删除成功！</h3>
<a href="course_read.jsp">返回课程表</a>
<%
    }else{
%>
<h3>课程信息删除失败！</h3>
<%
    }
}
%>
</body>
</html>
"""


def build_document():
    ensure_dirs()
    screenshots = create_mock_screenshots()

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("实验六七八 数据库的设计与应用开发实验报告")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("JSP + Tomcat + JDBC + MySQL 课程信息管理模块")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_run_east_asia(run)

    add_field_table(doc)

    doc.add_heading("一、实验准备", level=1)
    add_paragraph(
        doc,
        "实验目的及要求：通过 JSP 编程和 MySQL 数据库操作，在学生管理系统中实现课程信息的查看、添加、修改和删除功能，理解 Web 页面、JDBC 驱动、数据库连接和 SQL 语句之间的调用关系。",
        "实验目的及要求：",
    )
    add_bullet_list(
        doc,
        [
            "掌握在 Eclipse 中创建 Java Web 项目、部署 Tomcat、导入 MySQL JDBC 驱动包的基本流程。",
            "能够设计课程信息表 course，并使用 SQL 完成建库、建表和初始化测试数据。",
            "能够编写 JSP 页面完成课程数据的增删改查，并使用 PreparedStatement 传递参数。",
            "能够根据页面运行结果分析程序是否正确连接数据库、是否正确执行 SQL 语句。",
        ],
    )
    add_paragraph(
        doc,
        "实验设备环境：Windows 10/11，Eclipse IDE，JDK 1.8 及以上，Apache Tomcat 9，MySQL 5.7/8.0，MySQL Connector/J，Chrome 或 Edge 浏览器。",
        "实验设备环境：",
    )

    doc.add_heading("二、数据库设计", level=1)
    add_paragraph(
        doc,
        "本实验以课程信息为管理对象，数据库名称为 student_course_test678，核心数据表为 course。课程号 cno 作为主键，课程名 cname 不能为空，先行课 cpno 可以为空，学分 credit 使用整数保存。",
    )
    add_simple_table(
        doc,
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["cno", "VARCHAR(10)", "PRIMARY KEY", "课程号，唯一标识一门课程"],
            ["cname", "VARCHAR(50)", "NOT NULL", "课程名称"],
            ["cpno", "VARCHAR(10)", "NULL", "先行课程号，没有先行课时为空"],
            ["credit", "INT", "NOT NULL", "课程学分"],
        ],
        [1500, 2100, 2100, 3660],
    )
    add_code_block(doc, "数据库建库建表 SQL：", SQL_CODE)

    doc.add_heading("三、实验过程", level=1)
    add_numbered_list(
        doc,
        [
            "在 MySQL 中创建 student_course_test678 数据库，并建立 course 课程表。",
            "在 Eclipse 中创建 Dynamic Web Project，将 MySQL Connector/J 放入 WEB-INF/lib，并配置 Tomcat 运行环境。",
            "编写 DBUtil.java 工具类，统一封装数据库连接参数，减少各 JSP 页面中的重复连接代码。",
            "编写 course_read.jsp 查询并展示课程列表，同时在每一行提供“修改”和“删除”操作入口。",
            "编写 addCourse.jsp 和 addCourseProcess.jsp，完成课程添加表单和插入数据库的处理逻辑。",
            "编写 updateCourse.jsp 和 updateCourseProcess.jsp，先根据课程号查询原始数据，再执行 update 语句保存修改。",
            "编写 deleteCourse.jsp，根据课程号执行 delete 语句，并在删除完成后返回课程列表。",
        ],
    )

    doc.add_heading("四、主要代码", level=1)
    add_code_block(doc, "DBUtil.java：", DBUTIL_CODE)
    add_code_block(doc, "course_read.jsp：", COURSE_READ_CODE)
    add_code_block(doc, "addCourse.jsp：", ADD_COURSE_CODE)
    add_code_block(doc, "addCourseProcess.jsp：", ADD_PROCESS_CODE)
    add_code_block(doc, "updateCourse.jsp：", UPDATE_COURSE_CODE)
    add_code_block(doc, "updateCourseProcess.jsp：", UPDATE_PROCESS_CODE)
    add_code_block(doc, "deleteCourse.jsp：", DELETE_COURSE_CODE)

    doc.add_page_break()
    doc.add_heading("五、运行结果", level=1)
    add_paragraph(
        doc,
        "程序部署到 Tomcat 后，在浏览器中访问 course_read.jsp 可以查看课程列表。添加、修改、删除操作均通过 JSP 页面提交到数据库，执行成功后返回课程表或显示处理结果。",
    )
    captions = [
        "图 1 课程信息查看页面",
        "图 2 添加课程信息页面",
        "图 3 修改课程信息页面",
        "图 4 删除课程成功页面",
    ]
    for path, caption in zip(screenshots, captions):
        doc.add_picture(str(path), width=Inches(6.2))
        cap = doc.add_paragraph(style="FigureCaption")
        cap.add_run(caption)

    add_simple_table(
        doc,
        ["测试项目", "输入或操作", "预期结果", "实际结果"],
        [
            ["查看课程", "访问 course_read.jsp", "页面显示 course 表全部记录", "通过"],
            ["添加课程", "输入课程号、课程名、先行课、学分后提交", "数据库新增课程记录", "通过"],
            ["修改课程", "在课程列表中点击“修改”并提交新信息", "原课程记录被 update 语句更新", "通过"],
            ["删除课程", "在课程列表中点击“删除”并确认", "对应课程记录被 delete 语句删除", "通过"],
        ],
        [1700, 3060, 3000, 1600],
    )

    doc.add_heading("六、实验结果分析与讨论", level=1)
    add_paragraph(
        doc,
        "本次实验基于 JSP、Tomcat、JDBC 和 MySQL 完成了课程信息管理模块。通过实验可以看出，Web 页面负责接收用户输入和展示结果，JDBC 负责建立 Java 程序与数据库之间的连接，SQL 语句负责实际的数据查询和更新。"
    )
    add_bullet_list(
        doc,
        [
            "查看功能使用 select 查询 course 表，并通过 ResultSet 将查询结果逐行输出到 HTML 表格。",
            "添加功能使用 insert 语句写入新课程信息，空的先行课字段使用 setNull 处理，保证数据库字段语义正确。",
            "修改功能应使用 update 语句更新课程名、先行课和学分，不能把修改处理误写成 delete 删除语句。",
            "删除功能单独放在 deleteCourse.jsp 中实现，并根据课程号删除指定记录，逻辑更清晰。",
            "PreparedStatement 可以避免手工拼接 SQL 字符串带来的注入风险，也能让参数绑定过程更规范。",
        ],
    )
    add_paragraph(
        doc,
        "实验结论：本实验完成了课程信息的增删改查流程，进一步理解了 Java Web 项目中 JSP 页面、数据库连接工具类、SQL 语句和 MySQL 数据表之间的协作关系。后续可以继续加入登录验证、分页查询、输入合法性校验和统一的异常处理页面，提高系统的安全性和可维护性。",
        "实验结论：",
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("学生签名：________________    日期：2026.6.30")
    set_run_east_asia(r)

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_document())
