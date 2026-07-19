from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUT_DIR / "故宫防御系统科普视频作品说明报告.docx"
IMAGE_EFFECT = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-afbf9e58-b2be-4098-81bc-19bbcf2f02d5.png")
IMAGE_PROCESS = Path(r"C:\Users\Lenovo\AppData\Local\Temp\codex-clipboard-95d160cc-0d0d-4e5e-a2a7-79e346a914e4.png")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
BLACK = RGBColor(0, 0, 0)
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), "Calibri")
    r_pr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def para_format(paragraph, before=0, after=6, line=1.1, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_body(doc, text):
    p = doc.add_paragraph()
    para_format(p, after=6, line=1.1)
    run = p.add_run(text)
    set_run_font(run, 11, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        para_format(p, before=16, after=8, line=1.1)
        size = 16
    else:
        para_format(p, before=10, after=5, line=1.1)
        size = 13
    run = p.add_run(text)
    set_run_font(run, size, bold=True, color=BLUE)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    para_format(p, after=6, line=1.167)
    run = p.add_run(text)
    set_run_font(run, 11, color=BLACK)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    para_format(p, before=4, after=8, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, 9.5, color=MUTED)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=BLACK, size=10.5, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.text = ""
    para_format(p, after=0, line=1.1, align=align)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(total * 1440)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def metadata_table(doc):
    rows = [
        ("作品名称", "故宫防御系统科普视频"),
        ("作品类型", "科普类短视频 / AI 辅助影像创作"),
        ("主题定位", "讲解故宫古代防御体系与现代安全守护的结合"),
        ("小组成员", "姬鹏、伍文琪、赖鹏臣"),
        ("完成时间", "2026年6月"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1.35, 5.15])
    set_table_borders(table)
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        shade_cell(left, LIGHT_GRAY)
        set_cell_text(left, label, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(right, value)
    return table


def tools_table(doc):
    data = [
        ("AI 文案工具", "用于整理故宫防御系统的讲解逻辑，提炼城墙、城门、角楼、护城河、巡防与消防守护等知识点。"),
        ("AI 图像 / 视频生成工具", "用于生成或优化故宫场景、守卫、防御设施、现代消防守护等视觉画面。"),
        ("视频剪辑工具", "用于完成画面排序、转场、字幕、配音节奏和成片导出。"),
        ("字幕与配音工具", "用于制作讲解字幕，增强科普内容的可理解性。"),
        ("开放素材与字体", "使用可公开获取的参考资料、开放字体或免版权素材辅助排版与字幕呈现。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [1.55, 4.95])
    set_table_borders(table)
    for idx, label in enumerate(("工具 / 组件", "使用方式")):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], label, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for name, desc in data:
        cells = table.add_row().cells
        set_cell_text(cells[0], name, bold=True)
        set_cell_text(cells[1], desc)
    return table


def division_table(doc):
    data = [
        ("姬鹏", "负责主题确定、资料筛选、视频整体结构设计，以及故宫防御系统知识点整理。"),
        ("伍文琪", "负责画面素材收集、AI 画面生成提示词调整、字幕校对和作品效果截图整理。"),
        ("赖鹏臣", "负责剪辑合成、配音与字幕节奏调整、过程材料整理和最终文档排版。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, [1.2, 5.3])
    set_table_borders(table)
    for idx, label in enumerate(("成员", "具体分工")):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], label, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for member, work in data:
        cells = table.add_row().cells
        set_cell_text(cells[0], member, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], work)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167


def main():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    para_format(title, after=4, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    title_run = title.add_run("故宫防御系统科普视频作品说明报告")
    set_run_font(title_run, 23, bold=True, color=BLACK)

    subtitle = doc.add_paragraph()
    para_format(subtitle, after=18, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle_run = subtitle.add_run("从古代城防体系到现代科技守护")
    set_run_font(subtitle_run, 13, color=MUTED)

    metadata_table(doc)

    add_heading(doc, "（一）作品主要内容")
    add_body(
        doc,
        "本作品是一部以“故宫防御系统”为主题的科普视频。视频以故宫建筑群为视觉背景，围绕宫城防御、人员巡守、空间布局和现代安全保护展开讲解，说明故宫不仅是一座宫殿建筑群，也是一套兼具礼制、秩序和安全功能的城市级防御系统。"
    )
    add_body(
        doc,
        "作品主要介绍城墙、宫门、角楼、护城河、瞭望与巡防等传统防御元素，并进一步展示现代消防、监控预警、应急处置等科技手段如何参与文物保护。视频希望让观众理解：故宫的“防御”不只是抵御外敌，也包括防火、防盗、防灾和保护文化遗产。"
    )

    add_heading(doc, "（二）设计思路和重难点")
    add_body(
        doc,
        "设计上采用“古代防御设施 + 现代科技守护”的叙事结构。前半部分用角楼、宫墙、旗帜、守卫等画面建立历史感；后半部分用消防人员、应急设备和现代防护画面表现当代科技保护，使视频形成从历史到现实的递进关系。"
    )
    add_bullet(doc, "科普准确性：需要把“防御系统”讲清楚，避免只停留在画面好看，而忽略城墙、城门、角楼、护城河等结构的实际作用。")
    add_bullet(doc, "画面统一性：古建筑画面、守卫画面和现代消防画面风格差异较大，需要通过色调、字幕和剪辑节奏保持统一。")
    add_bullet(doc, "AI 画面控制：生成画面时容易出现建筑细节不准确、人物比例异常、服装时代混乱等问题，需要多次修改提示词并筛选可用片段。")
    add_bullet(doc, "科普节奏：视频需要在较短时间内讲清楚重点，因此文案要简明，字幕要突出关键词，避免信息过密。")

    add_heading(doc, "（三）使用开源组件和AI工具")
    add_body(
        doc,
        "本作品主要使用 AI 工具辅助完成文案梳理、画面生成、字幕润色和成片剪辑。开源或开放资源主要用于字体、参考资料整理和基础素材处理，没有依赖复杂的软件代码开发。"
    )
    tools_table(doc)

    add_heading(doc, "（四）作品效果截图")
    add_body(
        doc,
        "下图展示了视频中“现代科技守护古建筑”的核心画面：消防人员在故宫场景中执行应急守护任务，体现传统文化遗产保护与现代安全技术的结合。"
    )
    if IMAGE_EFFECT.exists():
        doc.add_picture(str(IMAGE_EFFECT), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "图 1 作品效果截图：现代消防力量守护故宫安全")
    else:
        add_body(doc, "（此处预留作品效果截图位置）")

    add_heading(doc, "（五）作品迭代过程中的过程性材料截图")
    add_body(
        doc,
        "迭代过程中重点调整了“古代防御系统”的视觉表达，例如角楼、瞭望台、旗帜、宫墙与远景环境。下图可作为分镜和画面筛选过程中的材料，用于表现古代防御设施的空间感和历史氛围。"
    )
    if IMAGE_PROCESS.exists():
        doc.add_picture(str(IMAGE_PROCESS), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "图 2 过程性材料截图：角楼与瞭望防御画面筛选")
    else:
        add_body(doc, "（此处预留过程性材料截图位置）")
    add_body(
        doc,
        "主要迭代流程为：确定科普主题；整理故宫防御系统资料；编写视频脚本；生成古代守卫、角楼、宫墙和现代消防画面；筛选稳定画面；加入字幕、配音和转场；最后检查科普信息是否清楚、画面是否连贯。"
    )

    add_heading(doc, "（六）作品具体分工")
    division_table(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("故宫防御系统科普视频作品说明报告")
    set_run_font(footer_run, 9, color=MUTED)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
